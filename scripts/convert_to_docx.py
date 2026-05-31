#!/usr/bin/env python3
"""
Converte relatórios Markdown para .docx e salva na pasta Downloads.
"""
import os
import sys
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Forcar UTF-8 no stdout para evitar erros de encoding no Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

DOWNLOADS_DIR = os.path.expanduser("~\\Downloads")

def md_to_docx(md_path, docx_name=None):
    """Converte um arquivo .md para .docx com formatação básica."""
    if docx_name is None:
        docx_name = os.path.splitext(os.path.basename(md_path))[0] + ".docx"
    
    docx_path = os.path.join(DOWNLOADS_DIR, docx_name)
    
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    
    doc = Document()
    
    # Configurar fonte padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = md_content.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Headers
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:], level=4)
        elif line.startswith("---"):
            # Horizontal rule
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            run.font.size = Pt(8)
        elif line.startswith("| "):
            # Table row - collect all rows
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Process table
            if len(table_lines) >= 2:
                rows_data = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split("|")[1:-1]]
                    rows_data.append(cells)
                
                # Skip separator row (|---|)
                if rows_data and all("---" in c for c in rows_data[1]):
                    rows_data.pop(1)
                
                if rows_data:
                    num_cols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=len(rows_data), cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                # Bold for header row
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
            continue
        elif line.strip() == "":
            # Empty line
            doc.add_paragraph()
        elif line.startswith("- ") or line.startswith("* "):
            # List item
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            # Numbered list
            p = doc.add_paragraph(line, style='List Number')
        elif line.startswith("**") and line.endswith("**"):
            # Bold line
            p = doc.add_paragraph()
            run = p.add_run(line.strip("**"))
            run.bold = True
        else:
            # Normal paragraph - handle inline formatting
            p = doc.add_paragraph()
            # Simple bold handling
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        i += 1
    
    doc.save(docx_path)
    return docx_path

def main():
    files = [
        {
            "md": "audits/CircleUSDCBridge/final_report.md",
            "docx": "Relatorio_Auditoria_Circle_CCTP_V2.docx"
        },
        {
            "md": "audits/Polymarket/final_report.md",
            "docx": "Relatorio_Auditoria_Polymarket_CTF_Exchange_V2.docx"
        },
        {
            "md": "audits/LayerZero/RELATORIO_FINAL.md",
            "docx": "Relatorio_Auditoria_LayerZero_V2.docx"
        },
        {
            "md": "knowledge_base/templates/immunefi_report_template.md",
            "docx": "Modelo_Relatorio_Bug_Bounty_Immunefi.docx"
        }
    ]
    
    print("=" * 60)
    print("🛡️  DeFi Security Workspace - Conversão para DOCX")
    print("=" * 60)
    
    for f in files:
        md_path = os.path.join(os.getcwd(), f["md"])
        if os.path.exists(md_path):
            print(f"\n📄 Convertendo: {f['md']}")
            docx_path = md_to_docx(md_path, f["docx"])
            print(f"   ✅ Salvo em: {docx_path}")
        else:
            print(f"\n❌ Arquivo não encontrado: {md_path}")
    
    print("\n" + "=" * 60)
    print("✅ Conversão concluída! Arquivos salvos em:")
    print(f"   {DOWNLOADS_DIR}")
    print("=" * 60)

if __name__ == "__main__":
    main()
